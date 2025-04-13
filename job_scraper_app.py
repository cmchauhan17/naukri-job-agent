import streamlit as st
from docx import Document
import io

# Function to read .docx file and extract text
def read_docx(file):
    doc = Document(file)
    text = ''
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

# Function to extract skills from the job description (basic keyword matching for now)
def extract_skills_from_job_desc(job_description):
    # Define a list of skills to look for
    skills = []
    keywords = ["Python", "Java", "SQL", "Leadership", "Teamwork", "Project Management"]
    
    for keyword in keywords:
        if keyword.lower() in job_description.lower():
            skills.append(keyword)
    
    return skills

# Function to tailor the resume based on job description
def tailor_resume(resume_text, job_description):
    tailored_doc = Document()

    # Split the resume text into paragraphs
    resume_paragraphs = resume_text.split('\n')

    # Extract skills based on the job description
    skills_to_add = extract_skills_from_job_desc(job_description)

    tailored_doc.add_paragraph("Resume Tailored for Job Description")

    # Go through each paragraph in the resume
    for para in resume_paragraphs:
        if "Skills" in para:  # Look for a "Skills" section in the resume
            tailored_doc.add_paragraph("Skills: " + ", ".join(skills_to_add))
        else:
            tailored_doc.add_paragraph(para)

    return tailored_doc

# Function to convert docx document to byte stream for download
def convert_docx_to_bytes(doc):
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# Streamlit UI
st.title("AI Resume Tailoring")

# Upload Resume
uploaded_resume = st.file_uploader("Upload your Resume (.docx)", type="docx")
job_description = st.text_area("Enter Job Description")

if uploaded_resume and job_description:
    # Read and display the uploaded resume
    resume_text = read_docx(uploaded_resume)
    st.text_area("Your Resume Text", value=resume_text, height=300)

    if st.button("Tailor Resume"):
        # Call the function to tailor the resume
        tailored_resume = tailor_resume(resume_text, job_description)

        # Convert the tailored resume to bytes for downloading
        tailored_resume_bytes = convert_docx_to_bytes(tailored_resume)

        # Provide the download link for the tailored resume
        st.download_button(
            label="Download Tailored Resume",
            data=tailored_resume_bytes,
            file_name="tailored_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


pip install transformers
pip install torch

from transformers import pipeline

# Load pre-trained model for text generation (T5 or GPT-2 for simplicity)
summarizer = pipeline("summarization", model="t5-small", tokenizer="t5-small")

# Function to rewrite resume content based on job description
def rewrite_text_for_job(description_text, resume_section):
    # Combine job description and resume section to guide the model
    input_text = f"Rewrite the following resume section based on the job description: {description_text}. Resume Section: {resume_section}"
    
    # Generate the rewritten text using the model
    rewritten_text = summarizer(input_text, max_length=200, min_length=50, do_sample=False)[0]['summary_text']
    
    return rewritten_text

# Function to tailor the resume with AI-powered rewriting
def tailor_resume_with_ai(resume_text, job_description):
    tailored_doc = Document()

    # Split the resume text into paragraphs
    resume_paragraphs = resume_text.split('\n')

    # Process each section (this is a basic example, you can refine it further)
    for para in resume_paragraphs:
        if "Skills" in para:  # Look for a "Skills" section in the resume
            skills_to_add = extract_skills_from_job_desc(job_description)
            tailored_doc.add_paragraph("Skills: " + ", ".join(skills_to_add))
        elif "Experience" in para:  # Look for an "Experience" section in the resume
            rewritten_experience = rewrite_text_for_job(job_description, para)
            tailored_doc.add_paragraph("Experience: " + rewritten_experience)
        else:
            tailored_doc.add_paragraph(para)

    return tailored_doc

if uploaded_resume and job_description:
    # Read and display the uploaded resume
    resume_text = read_docx(uploaded_resume)
    st.text_area("Your Resume Text", value=resume_text, height=300)

    if st.button("Tailor Resume"):
        # Call the AI-powered tailoring function
        tailored_resume = tailor_resume_with_ai(resume_text, job_description)

        # Convert the tailored resume to bytes for downloading
        tailored_resume_bytes = convert_docx_to_bytes(tailored_resume)

        # Provide the download link for the tailored resume
        st.download_button(
            label="Download Tailored Resume",
            data=tailored_resume_bytes,
            file_name="tailored_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
